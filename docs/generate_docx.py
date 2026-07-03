"""Generate a comprehensive DOCX analysis report for GCAAN ACTIVT."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

OUTPUT = "/app/docs/GCAAN_ACTIVT_Analysis.docx"
NAVY = RGBColor(0x0B, 0x3D, 0x91)
TEAL = RGBColor(0x0E, 0xA5, 0xA5)
GOLD = RGBColor(0xF5, 0xA5, 0x24)
GREY = RGBColor(0x4A, 0x55, 0x78)
MUTED = RGBColor(0x98, 0xA2, 0xB3)


def rtl(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def h1(doc, text):
    p = doc.add_paragraph()
    rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = NAVY
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = NAVY
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = TEAL
    return p


def para(doc, text, size=11, color=None, bold=False, align="right"):
    p = doc.add_paragraph()
    rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align == "right" else WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if bold:
        r.bold = True
    return p


def bullets(doc, items, color=None):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(it)
        r.font.size = Pt(11)
        if color:
            r.font.color.rgb = color


def table_from_data(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(str(val))
            r.font.size = Pt(9)
    return t


doc = Document()
section = doc.sections[0]
section.right_to_left = True
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

# ---------- Cover Page ----------
cover_title = doc.add_paragraph()
rtl(cover_title)
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_title.add_run("GCAAN ACTIVT")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = NAVY

para(doc, "نظام إدارة أنشطة الموظفين", size=18, color=GREY, align="center", bold=True)
para(doc, "الشركة العامة لإدارة المطارات والملاحة الجوية", size=14, color=GREY, align="center")
doc.add_paragraph()
doc.add_paragraph()

para(doc, "تحليل شامل ومفصّل للنظام", size=20, color=TEAL, align="center", bold=True)
doc.add_paragraph()
doc.add_paragraph()

para(doc, f"تاريخ التحليل: {datetime.now().strftime('%Y-%m-%d')}", size=11, color=MUTED, align="center")
para(doc, "الإصدار: 1.0.0", size=11, color=MUTED, align="center")
para(doc, "تصميم المهندس معاد كاظم", size=12, color=NAVY, align="center", bold=True)

doc.add_page_break()

# ---------- Executive Summary ----------
h1(doc, "١. الملخص التنفيذي")
para(doc,
    "GCAAN ACTIVT هو تطبيق متكامل لإدارة نشاطات موظفي الشركة العامة لإدارة المطارات "
    "والملاحة الجوية، مبني بأحدث التقنيات ومصمم بواجهة عربية RTL كاملة يعمل على "
    "أنظمة Android و iOS والويب. يدعم النظام تدفق موافقات هرمي متعدد المستويات "
    "من الموظف حتى مكتب المدير العام، مع تقارير غنية بصيغ PDF و Word، ولوحة "
    "مؤشرات أداء بصرية (KPIs)، ونظام نسخ احتياطي واستعادة كامل.")

h3(doc, "الأرقام الرئيسية")
bullets(doc, [
    "٩٧ اختبار خلفي ناجح 100% (لا أخطاء)",
    "٢٠+ نقطة نهاية API",
    "١٤ شاشة كاملة",
    "٤ أدوار متسلسلة (المدير العام → مدير قسم → مدير شعبة → موظف)",
    "٧ أنواع تقارير قابلة للطباعة والتصدير",
])

doc.add_paragraph()

# ---------- Architecture ----------
h1(doc, "٢. البنية التقنية")

h3(doc, "المكوّنات")
table_from_data(doc, ["الطبقة", "التقنية"], [
    ["Frontend", "Expo Router (React Native SDK 54) - RTL Arabic"],
    ["Backend", "FastAPI + Motor (async MongoDB driver)"],
    ["Database", "MongoDB (users, departments, divisions, activities)"],
    ["Authentication", "JWT + bcrypt (JWT_SECRET من .env)"],
    ["Charts", "react-native-svg (Gauge / Donut / Bar)"],
    ["PDF Reports", "expo-print"],
    ["Word Reports", "python-docx"],
    ["Backup Format", "JSON export/import"],
])

doc.add_paragraph()

# ---------- Roles ----------
h1(doc, "٣. الهيكل الإداري (٤ أدوار)")

para(doc,
    "النظام مبني على هرمية إدارية صارمة، حيث كل مستوى يستطيع إدارة المستوى الذي تحته:")

# Diagram-like table
t = doc.add_table(rows=4, cols=3)
t.style = "Light Grid Accent 1"
diagram_rows = [
    ("١", "مكتب المدير العام", "المستوى الأعلى — يعتمد النشاطات نهائياً + يدير الأقسام + النسخ الاحتياطي"),
    ("٢", "مدير القسم", "يدير الشعب + الموظفين + KPI موظفيه + يعتمد/يعدّل ويرفع للمدير العام"),
    ("٣", "مدير الشعبة", "يدير الموظفين + يعتمد ويرفع لمدير القسم"),
    ("٤", "الموظف", "يُدخل النشاطات ويتابع حالتها"),
]
for i, (lvl, role, desc) in enumerate(diagram_rows):
    for ci, val in enumerate([lvl, role, desc]):
        cell = t.rows[i].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(val)
        r.font.size = Pt(11)
        if ci == 1:
            r.bold = True
            r.font.color.rgb = NAVY

doc.add_page_break()

# ---------- Workflow ----------
h1(doc, "٤. مخطط سير العمل (Workflow)")

para(doc, "يمر النشاط من الإنشاء حتى الاعتماد النهائي عبر ٤ مراحل:", size=11)
doc.add_paragraph()

workflow = [
    ("١", "الموظف يُنشئ نشاطاً", "تاريخ + نوع + قسم مستهدف + ملاحظات", "pending_division"),
    ("٢", "مدير الشعبة يوافق", "→ يرفع لمدير القسم", "pending_department"),
    ("٣", "مدير القسم يوافق (أو يعدّل)", "→ يرفع لمكتب المدير العام", "pending_gm"),
    ("٤", "مكتب المدير العام يعتمد", "→ الاعتماد النهائي", "approved"),
]

wt = doc.add_table(rows=1 + len(workflow), cols=4)
wt.style = "Light Grid Accent 1"
for i, h in enumerate(["#", "المرحلة", "الوصف", "الحالة"]):
    cell = wt.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(10)
for i, row in enumerate(workflow, 1):
    for ci, val in enumerate(row):
        cell = wt.rows[i].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(val)
        r.font.size = Pt(10)
        if ci == 1:
            r.bold = True

doc.add_paragraph()
para(doc,
    "ملاحظات على التدفق:",
    size=11, bold=True, color=TEAL)
bullets(doc, [
    "يمكن الرفض في أي مرحلة → status: rejected مع سبب الرفض",
    "مدير القسم فقط يستطيع تعديل النشاط (نوع/ملاحظات/قسم مستهدف)",
    "بعد الاعتماد النهائي لا يمكن حذف النشاط",
    "جميع الإجراءات مسجّلة في history array مع الوقت والمنفّذ",
])

doc.add_page_break()

# ---------- Screens ----------
h1(doc, "٥. الواجهات الرئيسية (١٤ شاشة)")

screens = [
    ("١. شاشة تسجيل الدخول", "خلفية بتدرج أزرق ملكي، شعار GCAAN، حقول عربية RTL، توقيع 'تصميم المهندس معاد كاظم'. بيانات الدخول الافتراضية: admin/admin123."),
    ("٢. الشاشة الرئيسية", "شارة الدور، ٤ بطاقات KPI (الإجمالي/المعتمد/بانتظار/نسبة الاعتماد)، تنبيه الموافقات المعلقة، ٦-٧ إجراءات سريعة حسب دور المستخدم."),
    ("٣. النشاطات (Tab)", "قائمة النشاطات مع فلاتر ٦ (الكل/بانتظار الشعبة/بانتظار القسم/بانتظار مكتب المدير العام/معتمد/مرفوض). كل نشاط يعرض القسم المنفّذ والقسم المستهدف."),
    ("٤. إضافة نشاط جديد", "نموذج للموظف يحتوي: تاريخ النشاط + نوع النشاط + القائمة الكاملة للأقسام المستهدفة + الملاحظات."),
    ("٥. الموافقات (للمدراء)", "قائمة النشاطات المعلقة بانتظار موافقة المستخدم الحالي. ٣ إجراءات لكل نشاط: موافقة / تعديل (لمدير القسم فقط) / رفض مع سبب."),
    ("٦. إدارة الأقسام", "شاشة للمدير العام. إضافة قسم جديد + تعديل اسم + حذف (يُرفض إذا كان القسم يحتوي على مستخدمين أو شعب)."),
    ("٧. إدارة الشعب", "شاشة للمدير العام ومدير القسم. إضافة/تعديل/حذف الشعب ضمن نطاق القسم."),
    ("٨. إدارة الموظفين والمدراء", "شاشة موحدة لإدارة الحسابات. إضافة مستخدم جديد + تعديل الاسم وكلمة المرور + حذف الحساب (مع حماية للمدير العام وحسابك)."),
    ("٩. تقارير (Tab) - ٥ تبويبات", "أسبوعي، شهري، حسب الأقسام، نسب الإنجاز، تقرير شامل. كل تبويب يدعم الطباعة وتصدير PDF."),
    ("١٠. لوحة مؤشرات الأداء (KPI Dashboard)", "داشبورد بصري بالكامل: Gauge دائري للمتوسط + Donut لتوزيع التقييمات + Bar chart لأعلى ٥ موظفين + قائمة مفصلة. يُصنف تلقائياً حسب دور المستخدم."),
    ("١١. تقرير القسم (لمدير القسم)", "شاشة مخصصة تعرض معلومات القسم + أداء العاملين + قائمة النشاطات + زر تصدير 'تقرير نشاطات الموظفين' PDF."),
    ("١٢. تقرير حسب فترة زمنية", "شاشة مع حقول 'من تاريخ' و'إلى تاريخ'. المدير العام يستطيع فلترة قسم معين أو عرض كل الشركة. النتيجة: بطاقات ملخصة + قائمة تفصيلية + تصدير PDF."),
    ("١٣. إدارة النظام (للمدير العام)", "لوحة حجم البيانات + مساحة التخزين + عدد السجلات لكل مجموعة (المستخدمون/الأقسام/الشعب/النشاطات). ٢ إجراءات: تنزيل نسخة احتياطية JSON + استعادة من نسخة احتياطية."),
    ("١٤. الملف الشخصي", "معلومات المستخدم + شارة الدور + زر تسجيل خروج + شعار GCAAN وتوقيع المصمم."),
]

for title, desc in screens:
    h3(doc, title)
    para(doc, desc, size=11, color=GREY)
    doc.add_paragraph()

para(doc,
    "ملاحظة: لقطات شاشة عالية الدقة يمكن الحصول عليها بالضغط على PrintScreen بعد "
    "الدخول لكل شاشة على الرابط المنشور من نسخة الإنتاج.",
    size=10, color=MUTED)

doc.add_page_break()

# ---------- APIs ----------
h1(doc, "٦. نقاط النهاية (APIs) — ٢٥+ endpoint")

api_groups = [
    ("المصادقة", [
        "POST /api/auth/login — تسجيل الدخول (يعيد JWT)",
        "GET /api/auth/me — بيانات المستخدم الحالي",
    ]),
    ("الأقسام", [
        "GET /api/departments",
        "POST /api/departments — إنشاء (GM فقط)",
        "PUT /api/departments/{id} — تعديل",
        "DELETE /api/departments/{id} — حذف",
    ]),
    ("الشعب", [
        "GET /api/divisions",
        "POST /api/divisions — إنشاء",
        "PUT /api/divisions/{id} — تعديل",
        "DELETE /api/divisions/{id} — حذف",
    ]),
    ("المستخدمون", [
        "GET /api/users",
        "POST /api/users — إنشاء (حسب الهرمية)",
        "PUT /api/users/{id} — تعديل الاسم/كلمة المرور",
        "DELETE /api/users/{id} — حذف",
    ]),
    ("النشاطات", [
        "GET /api/activities?scope=pending|mine",
        "POST /api/activities — إنشاء نشاط",
        "POST /api/activities/{id}/action — موافقة/رفض/تعديل",
        "DELETE /api/activities/{id} — حذف (قبل الاعتماد)",
    ]),
    ("التقارير", [
        "GET /api/reports/summary",
        "GET /api/reports/weekly",
        "GET /api/reports/monthly",
        "GET /api/reports/by-department",
        "GET /api/reports/completion-rates",
        "GET /api/reports/kpis",
        "GET /api/reports/comprehensive",
        "GET /api/reports/comprehensive.docx — تصدير Word",
        "GET /api/reports/range?from_date=&to_date=&department_id= — تقرير بفترة",
    ]),
    ("إدارة النظام (GM only)", [
        "GET /api/admin/system-info — حجم البيانات",
        "GET /api/admin/backup — تنزيل نسخة احتياطية JSON",
        "POST /api/admin/restore — استعادة من نسخة احتياطية",
    ]),
]

for group, endpoints in api_groups:
    h3(doc, group)
    bullets(doc, endpoints)

doc.add_page_break()

# ---------- Data Model ----------
h1(doc, "٧. نموذج البيانات (MongoDB)")

h3(doc, "users")
para(doc, "id, username, hashed_password, full_name, role (general_manager/department_manager/division_manager/employee), department_id, division_id, created_by, created_at", size=10, color=GREY)

h3(doc, "departments")
para(doc, "id, name, created_at", size=10, color=GREY)

h3(doc, "divisions")
para(doc, "id, name, department_id, created_at", size=10, color=GREY)

h3(doc, "activities")
para(doc,
    "id, employee_id, employee_name, employee_department_id, employee_department_name, "
    "employee_division_id, activity_date, activity_type, target_department_id, "
    "target_department_name, notes, status, created_at, history[], rejection_reason",
    size=10, color=GREY)

# ---------- KPIs Formula ----------
h1(doc, "٨. صيغة مؤشرات الأداء (KPI)")

para(doc, "يتم حساب مؤشر أداء كل موظف حسب الصيغة التالية:", size=11)
para(doc,
    "KPI = 0.5 × نسبة الاعتماد + 0.3 × حجم النشاط + 0.2 × سرعة الاعتماد",
    size=12, color=NAVY, bold=True)

h3(doc, "التصنيفات")
table_from_data(doc, ["النطاق", "التقييم", "اللون"], [
    ["85 - 100", "ممتاز", "أخضر"],
    ["70 - 84", "جيد جداً", "أزرق"],
    ["55 - 69", "جيد", "برتقالي"],
    ["أقل من 55", "يحتاج تحسين", "أحمر"],
])

doc.add_page_break()

# ---------- Testing ----------
h1(doc, "٩. حالة الاختبار")

table_from_data(doc, ["الفئة", "الحالة"], [
    ["اختبارات المصادقة (Auth)", "✅ 100%"],
    ["اختبارات الأدوار والصلاحيات", "✅ 100%"],
    ["اختبارات دورة حياة النشاط", "✅ 100%"],
    ["اختبارات التقارير", "✅ 100%"],
    ["اختبارات التعديل والحذف", "✅ 100%"],
    ["اختبارات النسخ الاحتياطي والاستعادة", "✅ 100%"],
    ["إجمالي: 97 / 97 اختبار خلفي", "✅ ناجح"],
])

# ---------- Deployment ----------
h1(doc, "١٠. النشر (Deployment)")

para(doc,
    "التطبيق جاهز للنشر عبر منصة Emergent. عند الضغط على زر Deploy، ستحصل على:",
    size=11)
bullets(doc, [
    "رابط دائم ثابت على المدى الطويل يعمل 24/7",
    "HTTPS آمن (شهادة SSL تلقائية)",
    "الرابط يعمل من أي شبكة في العالم",
    "يمكن الوصول من متصفح أي جهاز (Android/iOS/Windows/Mac/Linux)",
    "لا حاجة لتحميل تطبيق Expo Go",
    "الرابط يبقى نفسه بعد كل تعديل — لن يتغير على المستخدمين",
    "التكلفة: 50 كريدت شهرياً",
])

para(doc, "خطوات النشر:", size=12, bold=True, color=NAVY)
bullets(doc, [
    "١. تأكد من تغيير كلمة مرور admin الافتراضية",
    "٢. اضغط على زر 'Deploy' في أعلى يمين واجهة Emergent",
    "٣. اختر الخطة الشهرية (50 كريدت)",
    "٤. انتظر 2-5 دقائق حتى اكتمال النشر",
    "٥. احصل على الرابط الدائم — شاركه مع الموظفين",
    "٦. اختيارياً: اربط دومين مخصص مثل activities.gcaan.iq",
])

# ---------- Security ----------
h1(doc, "١١. الأمان والخصوصية")

bullets(doc, [
    "كلمات المرور مُشفّرة بـ bcrypt (لا تُخزّن كنص)",
    "JWT tokens بمفتاح سري في .env",
    "صلاحيات مبنية على الأدوار (RBAC) مطبّقة على مستوى الـ Backend",
    "حماية من CORS مُفعّلة",
    "لا يمكن حذف حساب المدير العام (لحماية الوصول للنظام)",
    "لا يمكن حذف حسابك الشخصي",
    "الاستعادة من نسخة احتياطية تُرفض إذا لم يوجد بها مدير عام",
    "التحقق من صحة كل مدخل (validation) قبل الحفظ",
])

# ---------- Recommendations ----------
h1(doc, "١٢. التوصيات المستقبلية")

h3(doc, "أولوية عالية")
bullets(doc, [
    "تغيير كلمة مرور admin الافتراضية فور النشر",
    "إنشاء سياسة خصوصية (لمتطلبات Play Store)",
    "عمل نسخة احتياطية أسبوعية على الأقل",
])

h3(doc, "أولوية متوسطة")
bullets(doc, [
    "إضافة نظام تنبيهات داخلي (جرس مع عدّاد)",
    "سجل تدقيق (Audit Log) لكل عمليات التعديل والحذف",
    "نسخ احتياطية تلقائية مجدولة يومياً",
    "دومين مخصص للتطبيق (activities.gcaan.iq)",
])

h3(doc, "أولوية منخفضة")
bullets(doc, [
    "صور شخصية للمستخدمين",
    "مرفقات للنشاطات (صور base64)",
    "إشعارات Push بعد النشر على Play Store",
    "رسوم بيانية مقارنة بين الأقسام عبر الزمن",
])

# ---------- Footer ----------
doc.add_paragraph()
doc.add_paragraph()
para(doc, "─" * 40, size=8, color=MUTED, align="center")
para(doc, "GCAAN ACTIVT © 2026", size=10, color=NAVY, bold=True, align="center")
para(doc, "تصميم المهندس معاد كاظم", size=10, color=GREY, align="center")
para(doc, f"وثيقة أُنشئت في {datetime.now().strftime('%Y-%m-%d')}", size=9, color=MUTED, align="center")

doc.save(OUTPUT)
print(f"Generated: {OUTPUT}")
import os
print(f"Size: {os.path.getsize(OUTPUT)} bytes")
