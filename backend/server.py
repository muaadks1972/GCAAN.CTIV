from fastapi import FastAPI, APIRouter, Depends, HTTPException, status, Query, Body
from fastapi.responses import StreamingResponse
from fastapi.security import OAuth2PasswordBearer
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import io
import json
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional, Literal
import uuid
from datetime import datetime, timedelta, timezone
from jose import JWTError, jwt
from passlib.context import CryptContext
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# ---------- Config ----------
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

JWT_SECRET = os.environ['JWT_SECRET']
JWT_ALG = 'HS256'
JWT_EXPIRE_MIN = 60 * 24 * 7  # 7 days

pwd_ctx = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/api/auth/login", auto_error=False)

app = FastAPI(title="GCAAN Activities API")
api = APIRouter(prefix="/api")

# ---------- Utils ----------
def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()

# ---------- Audit Log ----------
async def log_audit(
    user: Optional[dict],
    action: str,
    entity_type: str,
    entity_id: Optional[str] = None,
    details: Optional[dict] = None,
    status_str: str = "success",
):
    """Insert an audit log entry. Silently handles errors so it never breaks main flow."""
    try:
        entry = {
            "id": str(uuid.uuid4()),
            "at": now_iso(),
            "user_id": (user or {}).get("id"),
            "user_name": (user or {}).get("full_name") or (user or {}).get("username"),
            "user_role": (user or {}).get("role"),
            "action": action,
            "entity_type": entity_type,
            "entity_id": entity_id,
            "details": details or {},
            "status": status_str,
        }
        await db.audit_logs.insert_one(entry)
    except Exception as e:
        logging.warning(f"audit log failed: {e}")

def hash_password(p: str) -> str:
    return pwd_ctx.hash(p)

def verify_password(p: str, h: str) -> bool:
    try:
        return pwd_ctx.verify(p, h)
    except Exception:
        return False

def create_token(user_id: str, role: str) -> str:
    payload = {
        "sub": user_id,
        "role": role,
        "exp": datetime.now(timezone.utc) + timedelta(minutes=JWT_EXPIRE_MIN),
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALG)

async def get_current_user(token: Optional[str] = Depends(oauth2_scheme)):
    if not token:
        raise HTTPException(status_code=401, detail="غير مصرح")
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALG])
        user_id = payload.get("sub")
    except JWTError:
        raise HTTPException(status_code=401, detail="رمز غير صالح")
    user = await db.users.find_one({"id": user_id}, {"_id": 0, "hashed_password": 0})
    if not user:
        raise HTTPException(status_code=401, detail="المستخدم غير موجود")
    return user

def require_roles(*roles):
    async def checker(user=Depends(get_current_user)):
        if user["role"] not in roles:
            raise HTTPException(status_code=403, detail="صلاحية غير كافية")
        return user
    return checker

# ---------- Models ----------
Role = Literal["general_manager", "department_manager", "division_manager", "employee"]
ActivityStatus = Literal["pending_division", "pending_department", "pending_gm", "approved", "rejected"]

class LoginIn(BaseModel):
    username: str
    password: str

class UserCreate(BaseModel):
    username: str
    password: str
    full_name: str
    role: Role
    department_id: Optional[str] = None
    division_id: Optional[str] = None

class DepartmentCreate(BaseModel):
    name: str

class DepartmentUpdate(BaseModel):
    name: str

class DivisionCreate(BaseModel):
    name: str
    department_id: str

class DivisionUpdate(BaseModel):
    name: str

class UserUpdate(BaseModel):
    full_name: Optional[str] = None
    password: Optional[str] = None

class ActivityCreate(BaseModel):
    activity_date: str  # YYYY-MM-DD
    activity_type: str
    target_department_id: Optional[str] = None  # القسم الذي تم عليه النشاط
    target_department_name: Optional[str] = None
    notes: Optional[str] = ""

class ActivityAction(BaseModel):
    action: Literal["approve", "reject", "edit"]
    reason: Optional[str] = None
    # for edit
    activity_type: Optional[str] = None
    notes: Optional[str] = None
    target_department_name: Optional[str] = None

# ---------- Seed ----------
DEFAULT_DEPARTMENTS = [
    "قسم العمليات الجوية",
    "قسم الملاحة الجوية",
    "قسم الأمن والسلامة",
    "قسم الصيانة والهندسة",
    "قسم الشؤون الإدارية والمالية",
]

async def seed_initial_data():
    # Departments
    if await db.departments.count_documents({}) == 0:
        for name in DEFAULT_DEPARTMENTS:
            await db.departments.insert_one({
                "id": str(uuid.uuid4()),
                "name": name,
                "created_at": now_iso(),
            })
        logging.info("Seeded departments")

    # Admin user (general manager)
    admin = await db.users.find_one({"username": "admin"})
    if not admin:
        await db.users.insert_one({
            "id": str(uuid.uuid4()),
            "username": "admin",
            "hashed_password": hash_password("admin123"),
            "full_name": "المدير العام",
            "role": "general_manager",
            "department_id": None,
            "division_id": None,
            "created_by": None,
            "created_at": now_iso(),
        })
        logging.info("Seeded admin general manager (admin/admin123)")

    # Backfill employee_department_name in old activities
    dept_map = {}
    async for d in db.departments.find({}, {"_id": 0}):
        dept_map[d["id"]] = d.get("name")
    cursor = db.activities.find(
        {"$or": [{"employee_department_name": {"$exists": False}}, {"employee_department_name": None}]},
        {"_id": 0, "id": 1, "employee_department_id": 1},
    )
    async for a in cursor:
        name = dept_map.get(a.get("employee_department_id"))
        if name:
            await db.activities.update_one({"id": a["id"]}, {"$set": {"employee_department_name": name}})

@app.on_event("startup")
async def _startup():
    await seed_initial_data()

@app.on_event("shutdown")
async def _shutdown():
    client.close()

# ---------- Health ----------
@api.get("/")
async def root():
    return {"ok": True, "app": "GCAAN"}

# ---------- Auth ----------
@api.post("/auth/login")
async def login(payload: LoginIn):
    user = await db.users.find_one({"username": payload.username})
    if not user or not verify_password(payload.password, user["hashed_password"]):
        await log_audit(
            {"username": payload.username},
            action="login",
            entity_type="auth",
            details={"username": payload.username},
            status_str="failed",
        )
        raise HTTPException(status_code=400, detail="اسم المستخدم أو كلمة المرور غير صحيحة")
    token = create_token(user["id"], user["role"])
    await log_audit(user, action="login", entity_type="auth", entity_id=user["id"])
    return {
        "access_token": token,
        "token_type": "bearer",
        "user": {
            "id": user["id"],
            "username": user["username"],
            "full_name": user["full_name"],
            "role": user["role"],
            "department_id": user.get("department_id"),
            "division_id": user.get("division_id"),
        },
    }

@api.get("/auth/me")
async def me(user=Depends(get_current_user)):
    return user

# ---------- Departments ----------
@api.get("/departments")
async def list_departments(user=Depends(get_current_user)):
    docs = await db.departments.find({}, {"_id": 0}).to_list(500)
    return docs

@api.post("/departments")
async def create_department(payload: DepartmentCreate, user=Depends(require_roles("general_manager"))):
    doc = {"id": str(uuid.uuid4()), "name": payload.name, "created_at": now_iso()}
    await db.departments.insert_one(doc)
    doc.pop("_id", None)
    await log_audit(user, action="create", entity_type="department", entity_id=doc["id"], details={"name": doc["name"]})
    return doc

@api.put("/departments/{dept_id}")
async def update_department(dept_id: str, payload: DepartmentUpdate, user=Depends(require_roles("general_manager"))):
    name = (payload.name or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="اسم القسم مطلوب")
    result = await db.departments.update_one({"id": dept_id}, {"$set": {"name": name}})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="القسم غير موجود")
    # Update denormalized names in activities (both target and performing)
    await db.activities.update_many(
        {"target_department_id": dept_id},
        {"$set": {"target_department_name": name}},
    )
    await db.activities.update_many(
        {"employee_department_id": dept_id},
        {"$set": {"employee_department_name": name}},
    )
    doc = await db.departments.find_one({"id": dept_id}, {"_id": 0})
    await log_audit(user, action="update", entity_type="department", entity_id=dept_id, details={"name": name})
    return doc

@api.delete("/departments/{dept_id}")
async def delete_department(dept_id: str, user=Depends(require_roles("general_manager"))):
    if await db.users.count_documents({"department_id": dept_id}) > 0:
        raise HTTPException(status_code=400, detail="لا يمكن حذف قسم يحتوي على مستخدمين")
    if await db.divisions.count_documents({"department_id": dept_id}) > 0:
        raise HTTPException(status_code=400, detail="لا يمكن حذف قسم يحتوي على شعب. احذف الشعب أولاً")
    result = await db.departments.delete_one({"id": dept_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="القسم غير موجود")
    await log_audit(user, action="delete", entity_type="department", entity_id=dept_id)
    return {"ok": True}

# ---------- Divisions ----------
@api.get("/divisions")
async def list_divisions(department_id: Optional[str] = None, user=Depends(get_current_user)):
    q = {}
    if department_id:
        q["department_id"] = department_id
    docs = await db.divisions.find(q, {"_id": 0}).to_list(500)
    return docs

@api.post("/divisions")
async def create_division(payload: DivisionCreate, user=Depends(require_roles("general_manager", "department_manager"))):
    # Department manager can only create in own department
    if user["role"] == "department_manager" and user.get("department_id") != payload.department_id:
        raise HTTPException(status_code=403, detail="لا يمكنك إنشاء شعبة خارج قسمك")
    doc = {
        "id": str(uuid.uuid4()),
        "name": payload.name,
        "department_id": payload.department_id,
        "created_at": now_iso(),
    }
    await db.divisions.insert_one(doc)
    doc.pop("_id", None)
    await log_audit(user, action="create", entity_type="division", entity_id=doc["id"], details={"name": doc["name"], "department_id": doc["department_id"]})
    return doc

@api.put("/divisions/{div_id}")
async def update_division(div_id: str, payload: DivisionUpdate, user=Depends(require_roles("general_manager", "department_manager"))):
    div = await db.divisions.find_one({"id": div_id}, {"_id": 0})
    if not div:
        raise HTTPException(status_code=404, detail="الشعبة غير موجودة")
    if user["role"] == "department_manager" and div.get("department_id") != user.get("department_id"):
        raise HTTPException(status_code=403, detail="لا يمكنك تعديل شعبة خارج قسمك")
    name = (payload.name or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="اسم الشعبة مطلوب")
    await db.divisions.update_one({"id": div_id}, {"$set": {"name": name}})
    div["name"] = name
    await log_audit(user, action="update", entity_type="division", entity_id=div_id, details={"name": name})
    return div

@api.delete("/divisions/{div_id}")
async def delete_division(div_id: str, user=Depends(require_roles("general_manager", "department_manager"))):
    div = await db.divisions.find_one({"id": div_id}, {"_id": 0})
    if not div:
        raise HTTPException(status_code=404, detail="الشعبة غير موجودة")
    if user["role"] == "department_manager" and div.get("department_id") != user.get("department_id"):
        raise HTTPException(status_code=403, detail="لا يمكنك حذف شعبة خارج قسمك")
    if await db.users.count_documents({"division_id": div_id}) > 0:
        raise HTTPException(status_code=400, detail="لا يمكن حذف شعبة تحتوي على مستخدمين")
    await db.divisions.delete_one({"id": div_id})
    await log_audit(user, action="delete", entity_type="division", entity_id=div_id)
    return {"ok": True}


# ---------- Users ----------
@api.get("/users")
async def list_users(user=Depends(get_current_user)):
    q = {}
    if user["role"] == "general_manager":
        pass
    elif user["role"] == "department_manager":
        q["department_id"] = user.get("department_id")
    elif user["role"] == "division_manager":
        q["division_id"] = user.get("division_id")
    else:
        q["id"] = user["id"]  # employees see only themselves
    docs = await db.users.find(q, {"_id": 0, "hashed_password": 0}).to_list(1000)
    return docs

@api.post("/users")
async def create_user(payload: UserCreate, user=Depends(get_current_user)):
    # Role hierarchy checks
    role = payload.role
    if user["role"] == "general_manager":
        if role != "department_manager":
            raise HTTPException(status_code=403, detail="المدير العام يُنشئ مديري الأقسام فقط")
        if not payload.department_id:
            raise HTTPException(status_code=400, detail="يجب تحديد القسم")
    elif user["role"] == "department_manager":
        if role != "division_manager":
            raise HTTPException(status_code=403, detail="مدير القسم يُنشئ مديري الشعب فقط")
        if payload.department_id != user.get("department_id"):
            raise HTTPException(status_code=403, detail="خارج نطاق قسمك")
        if not payload.division_id:
            raise HTTPException(status_code=400, detail="يجب تحديد الشعبة")
    elif user["role"] == "division_manager":
        if role != "employee":
            raise HTTPException(status_code=403, detail="مدير الشعبة يُنشئ الموظفين فقط")
        payload.department_id = user.get("department_id")
        payload.division_id = user.get("division_id")
    else:
        raise HTTPException(status_code=403, detail="غير مصرح")

    if await db.users.find_one({"username": payload.username}):
        raise HTTPException(status_code=400, detail="اسم المستخدم مستخدم مسبقاً")

    doc = {
        "id": str(uuid.uuid4()),
        "username": payload.username,
        "hashed_password": hash_password(payload.password),
        "full_name": payload.full_name,
        "role": role,
        "department_id": payload.department_id,
        "division_id": payload.division_id,
        "created_by": user["id"],
        "created_at": now_iso(),
    }
    await db.users.insert_one(doc)
    await log_audit(user, action="create", entity_type="user", entity_id=doc["id"], details={"username": doc["username"], "role": doc["role"], "full_name": doc["full_name"]})
    return {k: v for k, v in doc.items() if k not in ("_id", "hashed_password")}

@api.put("/users/{user_id}")
async def update_user(user_id: str, payload: UserUpdate, user=Depends(get_current_user)):
    target = await db.users.find_one({"id": user_id})
    if not target:
        raise HTTPException(status_code=404, detail="المستخدم غير موجود")

    # Permission: GM edits department_managers; dept_mgr edits division_managers + employees in own dept;
    # div_mgr edits employees in own division; anyone can edit themselves (name+password).
    allowed = False
    if user["id"] == user_id:
        allowed = True
    elif user["role"] == "general_manager" and target.get("role") == "department_manager":
        allowed = True
    elif user["role"] == "department_manager" and target.get("department_id") == user.get("department_id") and target.get("role") in ("division_manager", "employee"):
        allowed = True
    elif user["role"] == "division_manager" and target.get("division_id") == user.get("division_id") and target.get("role") == "employee":
        allowed = True

    if not allowed:
        raise HTTPException(status_code=403, detail="لا تملك صلاحية تعديل هذا المستخدم")

    updates = {}
    if payload.full_name is not None:
        name = payload.full_name.strip()
        if not name:
            raise HTTPException(status_code=400, detail="الاسم لا يمكن أن يكون فارغاً")
        updates["full_name"] = name
    if payload.password is not None:
        if len(payload.password) < 4:
            raise HTTPException(status_code=400, detail="كلمة المرور يجب أن تكون 4 أحرف فأكثر")
        updates["hashed_password"] = hash_password(payload.password)

    if not updates:
        raise HTTPException(status_code=400, detail="لا توجد تعديلات")

    await db.users.update_one({"id": user_id}, {"$set": updates})

    # Update denormalized employee_name on their activities
    if "full_name" in updates:
        await db.activities.update_many({"employee_id": user_id}, {"$set": {"employee_name": updates["full_name"]}})

    updated = await db.users.find_one({"id": user_id}, {"_id": 0, "hashed_password": 0})
    await log_audit(user, action="update", entity_type="user", entity_id=user_id, details={"fields": list(updates.keys())})
    return updated

@api.delete("/users/{user_id}")
async def delete_user(user_id: str, user=Depends(get_current_user)):
    target = await db.users.find_one({"id": user_id})
    if not target:
        raise HTTPException(status_code=404, detail="المستخدم غير موجود")
    if target["role"] == "general_manager":
        raise HTTPException(status_code=403, detail="لا يمكن حذف المدير العام")
    if user["id"] == user_id:
        raise HTTPException(status_code=400, detail="لا يمكن حذف حسابك الشخصي")
    allowed = False
    if user["role"] == "general_manager" and target.get("role") == "department_manager":
        allowed = True
    elif user["role"] == "department_manager" and target.get("department_id") == user.get("department_id") and target.get("role") in ("division_manager", "employee"):
        allowed = True
    elif user["role"] == "division_manager" and target.get("division_id") == user.get("division_id") and target.get("role") == "employee":
        allowed = True
    if not allowed:
        raise HTTPException(status_code=403, detail="لا تملك صلاحية حذف هذا المستخدم")
    await db.users.delete_one({"id": user_id})
    await log_audit(user, action="delete", entity_type="user", entity_id=user_id, details={"username": target.get("username"), "role": target.get("role")})
    return {"ok": True}


# ---------- Activities ----------
def initial_activity_status(user_role: str) -> str:
    if user_role == "employee":
        return "pending_division"
    if user_role == "division_manager":
        return "pending_department"
    if user_role == "department_manager":
        return "pending_gm"
    return "approved"  # general_manager writing an activity is auto-approved (rare)

@api.post("/activities")
async def create_activity(payload: ActivityCreate, user=Depends(get_current_user)):
    dept_name = payload.target_department_name
    if payload.target_department_id and not dept_name:
        d = await db.departments.find_one({"id": payload.target_department_id}, {"_id": 0})
        if d:
            dept_name = d.get("name")
    # Performing department (department of the employee doing the activity)
    employee_dept_name = None
    if user.get("department_id"):
        emp_dept = await db.departments.find_one({"id": user["department_id"]}, {"_id": 0})
        if emp_dept:
            employee_dept_name = emp_dept.get("name")
    doc = {
        "id": str(uuid.uuid4()),
        "employee_id": user["id"],
        "employee_name": user["full_name"],
        "employee_department_id": user.get("department_id"),
        "employee_department_name": employee_dept_name,
        "employee_division_id": user.get("division_id"),
        "activity_date": payload.activity_date,
        "activity_type": payload.activity_type,
        "target_department_id": payload.target_department_id,
        "target_department_name": dept_name,
        "notes": payload.notes or "",
        "status": initial_activity_status(user["role"]),
        "created_at": now_iso(),
        "history": [
            {"by": user["id"], "by_name": user["full_name"], "action": "created", "at": now_iso()}
        ],
        "rejection_reason": None,
    }
    await db.activities.insert_one(doc)
    doc.pop("_id", None)
    await log_audit(user, action="create", entity_type="activity", entity_id=doc["id"], details={"activity_type": doc["activity_type"], "target_department": dept_name})
    return doc

@api.get("/activities")
async def list_activities(
    status_filter: Optional[str] = Query(None, alias="status"),
    scope: Optional[str] = Query(None),  # 'mine' | 'pending' | 'all'
    user=Depends(get_current_user),
):
    q = {}
    if status_filter:
        q["status"] = status_filter

    if scope == "mine" or user["role"] == "employee":
        q["employee_id"] = user["id"]
    elif user["role"] == "division_manager":
        q["employee_division_id"] = user.get("division_id")
        if scope == "pending":
            q["status"] = "pending_division"
    elif user["role"] == "department_manager":
        q["employee_department_id"] = user.get("department_id")
        if scope == "pending":
            q["status"] = "pending_department"
    elif user["role"] == "general_manager":
        if scope == "pending":
            q["status"] = "pending_gm"

    docs = await db.activities.find(q, {"_id": 0}).sort("created_at", -1).to_list(2000)
    return docs

@api.post("/activities/{activity_id}/action")
async def act_on_activity(activity_id: str, payload: ActivityAction, user=Depends(get_current_user)):
    activity = await db.activities.find_one({"id": activity_id}, {"_id": 0})
    if not activity:
        raise HTTPException(status_code=404, detail="النشاط غير موجود")

    role = user["role"]
    st = activity["status"]

    # Permission per role/status
    allowed = False
    next_status = st
    if role == "division_manager" and st == "pending_division" and activity.get("employee_division_id") == user.get("division_id"):
        allowed = True
        next_status = "pending_department" if payload.action == "approve" else st
    elif role == "department_manager" and st == "pending_department" and activity.get("employee_department_id") == user.get("department_id"):
        allowed = True
        next_status = "pending_gm" if payload.action == "approve" else st
    elif role == "general_manager" and st == "pending_gm":
        allowed = True
        next_status = "approved" if payload.action == "approve" else st

    if not allowed:
        raise HTTPException(status_code=403, detail="لا يمكنك تنفيذ هذا الإجراء الآن")

    updates = {}
    if payload.action == "reject":
        next_status = "rejected"
        updates["rejection_reason"] = payload.reason or ""
    if payload.action == "edit":
        # only department manager can edit per spec
        if role != "department_manager":
            raise HTTPException(status_code=403, detail="التعديل مسموح لمدير القسم فقط")
        if payload.activity_type is not None:
            updates["activity_type"] = payload.activity_type
        if payload.notes is not None:
            updates["notes"] = payload.notes
        if payload.target_department_name is not None:
            updates["target_department_name"] = payload.target_department_name
        # Editing doesn't advance status by itself unless action was 'edit' only
        next_status = st

    updates["status"] = next_status
    hist = activity.get("history", [])
    hist.append({
        "by": user["id"],
        "by_name": user["full_name"],
        "action": payload.action,
        "reason": payload.reason,
        "at": now_iso(),
    })
    updates["history"] = hist

    await db.activities.update_one({"id": activity_id}, {"$set": updates})
    activity.update(updates)
    await log_audit(user, action=payload.action, entity_type="activity", entity_id=activity_id, details={"new_status": next_status, "reason": payload.reason, "activity_type": activity.get("activity_type"), "employee_name": activity.get("employee_name")})
    return activity

@api.delete("/activities/{activity_id}")
async def delete_activity(activity_id: str, user=Depends(get_current_user)):
    activity = await db.activities.find_one({"id": activity_id}, {"_id": 0})
    if not activity:
        raise HTTPException(status_code=404, detail="غير موجود")
    if activity["employee_id"] != user["id"] and user["role"] != "general_manager":
        raise HTTPException(status_code=403, detail="غير مصرح")
    if activity["status"] not in ("pending_division", "rejected"):
        raise HTTPException(status_code=400, detail="لا يمكن حذف النشاط بعد اعتماده")
    await db.activities.delete_one({"id": activity_id})
    await log_audit(user, action="delete", entity_type="activity", entity_id=activity_id, details={"activity_type": activity.get("activity_type")})
    return {"ok": True}

# ---------- Reports ----------
def _filter_by_scope(user):
    """Scope filter for REPORTS.

    - employee → sees full DIVISION-level data (all colleagues in same division)
    - division_manager → sees own division
    - department_manager → sees own department
    - general_manager → sees everything
    """
    q = {}
    if user["role"] == "employee":
        q["employee_division_id"] = user.get("division_id")
    elif user["role"] == "division_manager":
        q["employee_division_id"] = user.get("division_id")
    elif user["role"] == "department_manager":
        q["employee_department_id"] = user.get("department_id")
    return q

@api.get("/reports/summary")
async def report_summary(user=Depends(get_current_user)):
    q = _filter_by_scope(user)
    total = await db.activities.count_documents(q)
    approved = await db.activities.count_documents({**q, "status": "approved"})
    rejected = await db.activities.count_documents({**q, "status": "rejected"})
    pending = total - approved - rejected
    approval_rate = round((approved / total) * 100, 1) if total else 0.0
    return {
        "total": total,
        "approved": approved,
        "rejected": rejected,
        "pending": pending,
        "approval_rate": approval_rate,
    }

@api.get("/reports/weekly")
async def report_weekly(user=Depends(get_current_user)):
    q = _filter_by_scope(user)
    since = (datetime.now(timezone.utc) - timedelta(days=7)).isoformat()
    q["created_at"] = {"$gte": since}
    docs = await db.activities.find(q, {"_id": 0}).to_list(5000)
    return {"period": "weekly", "count": len(docs), "activities": docs}

@api.get("/reports/monthly")
async def report_monthly(user=Depends(get_current_user)):
    q = _filter_by_scope(user)
    since = (datetime.now(timezone.utc) - timedelta(days=30)).isoformat()
    q["created_at"] = {"$gte": since}
    docs = await db.activities.find(q, {"_id": 0}).to_list(5000)
    return {"period": "monthly", "count": len(docs), "activities": docs}

@api.get("/reports/range")
async def report_range(
    from_date: str = Query(..., description="YYYY-MM-DD"),
    to_date: str = Query(..., description="YYYY-MM-DD"),
    department_id: Optional[str] = None,
    user=Depends(get_current_user),
):
    """Activities within a specific date range, optionally filtered by department (GM only for arbitrary dept)."""
    try:
        start = datetime.fromisoformat(from_date).replace(tzinfo=timezone.utc)
        end = (datetime.fromisoformat(to_date) + timedelta(days=1)).replace(tzinfo=timezone.utc)
    except Exception:
        raise HTTPException(status_code=400, detail="صيغة التاريخ غير صحيحة (استخدم YYYY-MM-DD)")
    if end <= start:
        raise HTTPException(status_code=400, detail="تاريخ النهاية يجب أن يكون بعد تاريخ البداية")

    q = _filter_by_scope(user)
    q["created_at"] = {"$gte": start.isoformat(), "$lt": end.isoformat()}

    # GM can additionally filter by a specific department; other roles are already scoped
    if department_id and user["role"] == "general_manager":
        q["employee_department_id"] = department_id

    docs = await db.activities.find(q, {"_id": 0}).sort("created_at", -1).to_list(10000)
    approved = sum(1 for a in docs if a.get("status") == "approved")
    rejected = sum(1 for a in docs if a.get("status") == "rejected")
    pending = len(docs) - approved - rejected

    # Group by department (performing)
    by_dept: dict = {}
    for a in docs:
        name = a.get("employee_department_name") or "بدون قسم"
        d = by_dept.setdefault(name, {"name": name, "total": 0, "approved": 0, "rejected": 0, "pending": 0})
        d["total"] += 1
        if a.get("status") == "approved":
            d["approved"] += 1
        elif a.get("status") == "rejected":
            d["rejected"] += 1
        else:
            d["pending"] += 1
    by_dept_list = sorted(by_dept.values(), key=lambda x: x["total"], reverse=True)

    return {
        "from_date": from_date,
        "to_date": to_date,
        "department_id": department_id,
        "count": len(docs),
        "approved": approved,
        "rejected": rejected,
        "pending": pending,
        "by_department": by_dept_list,
        "activities": docs,
    }


@api.get("/reports/by-department")
async def report_by_department(user=Depends(get_current_user)):
    q = _filter_by_scope(user)
    dept_filter: dict = {}
    if user["role"] in ("employee", "division_manager", "department_manager"):
        dept_id = user.get("department_id")
        if dept_id:
            dept_filter["id"] = dept_id
    departments = await db.departments.find(dept_filter, {"_id": 0}).to_list(500)
    # Single aggregation instead of N*4 count_documents
    pipeline = [
        {"$match": q},
        {"$group": {
            "_id": "$employee_department_id",
            "total": {"$sum": 1},
            "approved": {"$sum": {"$cond": [{"$eq": ["$status", "approved"]}, 1, 0]}},
            "rejected": {"$sum": {"$cond": [{"$eq": ["$status", "rejected"]}, 1, 0]}},
        }},
    ]
    stats: dict = {}
    async for row in db.activities.aggregate(pipeline):
        stats[row["_id"]] = row
    result = []
    for d in departments:
        s = stats.get(d["id"], {"total": 0, "approved": 0, "rejected": 0})
        total = s["total"]
        approved = s["approved"]
        rejected = s["rejected"]
        pending = total - approved - rejected
        result.append({
            "department_id": d["id"],
            "department_name": d["name"],
            "total": total,
            "approved": approved,
            "rejected": rejected,
            "pending": pending,
            "approval_rate": round((approved / total) * 100, 1) if total else 0.0,
        })
    return result

@api.get("/reports/completion-rates")
async def report_completion_rates(user=Depends(get_current_user)):
    """Per-employee completion (approval) rate."""
    q = _filter_by_scope(user)
    users_q = {"role": "employee"}
    if user["role"] == "division_manager":
        users_q["division_id"] = user.get("division_id")
    elif user["role"] == "department_manager":
        users_q["department_id"] = user.get("department_id")
    elif user["role"] == "employee":
        users_q["division_id"] = user.get("division_id")
    employees = await db.users.find(users_q, {"_id": 0, "hashed_password": 0}).to_list(2000)

    pipeline = [
        {"$match": q},
        {"$group": {
            "_id": "$employee_id",
            "total": {"$sum": 1},
            "approved": {"$sum": {"$cond": [{"$eq": ["$status", "approved"]}, 1, 0]}},
            "rejected": {"$sum": {"$cond": [{"$eq": ["$status", "rejected"]}, 1, 0]}},
        }},
    ]
    stats: dict = {}
    async for row in db.activities.aggregate(pipeline):
        stats[row["_id"]] = row

    result = []
    for e in employees:
        s = stats.get(e["id"], {"total": 0, "approved": 0, "rejected": 0})
        total = s["total"]
        approved = s["approved"]
        rejected = s["rejected"]
        result.append({
            "employee_id": e["id"],
            "employee_name": e["full_name"],
            "total": total,
            "approved": approved,
            "rejected": rejected,
            "completion_rate": round((approved / total) * 100, 1) if total else 0.0,
        })
    result.sort(key=lambda x: x["completion_rate"], reverse=True)
    return result

@api.get("/reports/kpis")
async def report_kpis(user=Depends(get_current_user)):
    """KPI dashboard per employee: activity count, approval rate, avg turnaround (days)."""
    q = _filter_by_scope(user)
    users_q = {"role": "employee"}
    if user["role"] == "division_manager":
        users_q["division_id"] = user.get("division_id")
    elif user["role"] == "department_manager":
        users_q["department_id"] = user.get("department_id")
    elif user["role"] == "employee":
        users_q["division_id"] = user.get("division_id")
    employees = await db.users.find(users_q, {"_id": 0, "hashed_password": 0}).to_list(2000)

    # Aggregate counts by employee
    pipeline = [
        {"$match": q},
        {"$group": {
            "_id": "$employee_id",
            "total": {"$sum": 1},
            "approved": {"$sum": {"$cond": [{"$eq": ["$status", "approved"]}, 1, 0]}},
        }},
    ]
    counts: dict = {}
    async for row in db.activities.aggregate(pipeline):
        counts[row["_id"]] = row

    # Fetch approved activities in one batch (with history) for turnaround
    approved_by_emp: dict = {}
    approved_docs = await db.activities.find({**q, "status": "approved"}, {"_id": 0}).to_list(5000)
    for a in approved_docs:
        approved_by_emp.setdefault(a["employee_id"], []).append(a)

    kpis = []
    for e in employees:
        c = counts.get(e["id"], {"total": 0, "approved": 0})
        total = c["total"]
        approved = c["approved"]
        approval_rate = round((approved / total) * 100, 1) if total else 0.0
        turnarounds = []
        for a in approved_by_emp.get(e["id"], []):
            try:
                created = datetime.fromisoformat(a["created_at"])
                hist = a.get("history", [])
                approvals = [h for h in hist if h.get("action") == "approve"]
                if approvals:
                    finished = datetime.fromisoformat(approvals[-1]["at"])
                    turnarounds.append((finished - created).total_seconds() / 86400.0)
            except Exception:
                continue
        avg_turnaround = round(sum(turnarounds) / len(turnarounds), 2) if turnarounds else 0.0
        vol_score = min(total, 20) / 20 * 100
        speed_score = 100 if not turnarounds else max(0, 100 - avg_turnaround * 10)
        kpi_score = round(0.5 * approval_rate + 0.3 * vol_score + 0.2 * speed_score, 1)
        rating = "ممتاز" if kpi_score >= 85 else "جيد جداً" if kpi_score >= 70 else "جيد" if kpi_score >= 55 else "يحتاج تحسين"
        kpis.append({
            "employee_id": e["id"],
            "employee_name": e["full_name"],
            "total_activities": total,
            "approved": approved,
            "approval_rate": approval_rate,
            "avg_turnaround_days": avg_turnaround,
            "kpi_score": kpi_score,
            "rating": rating,
        })
    kpis.sort(key=lambda x: x["kpi_score"], reverse=True)
    return kpis

# ---------- Comprehensive Report + DOCX Export ----------
def _set_rtl(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)


async def _build_comprehensive_data(user):
    q = _filter_by_scope(user)
    activities = await db.activities.find(q, {"_id": 0}).sort("created_at", -1).to_list(5000)
    # aggregate per (employee_id, target_department_id)
    per_emp: dict = {}
    for a in activities:
        eid = a.get("employee_id")
        if eid not in per_emp:
            per_emp[eid] = {
                "employee_id": eid,
                "employee_name": a.get("employee_name") or "-",
                "employee_department_name": a.get("employee_department_name") or "-",
                "total": 0,
                "approved": 0,
                "targets": {},
            }
        rec = per_emp[eid]
        rec["total"] += 1
        if a.get("status") == "approved":
            rec["approved"] += 1
        tname = a.get("target_department_name") or "-"
        rec["targets"][tname] = rec["targets"].get(tname, 0) + 1
    result = []
    for r in per_emp.values():
        targets_list = [{"name": n, "count": c} for n, c in r["targets"].items()]
        targets_list.sort(key=lambda x: x["count"], reverse=True)
        r["targets"] = targets_list
        result.append(r)
    result.sort(key=lambda x: x["total"], reverse=True)
    return result


@api.get("/reports/comprehensive")
async def report_comprehensive(user=Depends(get_current_user)):
    return await _build_comprehensive_data(user)


@api.get("/reports/comprehensive.docx")
async def report_comprehensive_docx(user=Depends(get_current_user)):
    data = await _build_comprehensive_data(user)
    doc = Document()

    section = doc.sections[0]
    section.right_to_left = True

    # Title
    title = doc.add_paragraph()
    _set_rtl(title)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = title.add_run("تقرير شامل بنشاطات الموظفين")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x0B, 0x3D, 0x91)

    sub = doc.add_paragraph()
    _set_rtl(sub)
    sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    s = sub.add_run("الشركة العامة لإدارة المطارات والملاحة الجوية")
    s.font.size = Pt(11)
    s.font.color.rgb = RGBColor(0x4A, 0x55, 0x78)

    date_p = doc.add_paragraph()
    _set_rtl(date_p)
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    d = date_p.add_run(datetime.now(timezone.utc).strftime("%Y-%m-%d"))
    d.font.size = Pt(10)
    d.font.color.rgb = RGBColor(0x98, 0xA2, 0xB3)

    doc.add_paragraph()

    # Table headers
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "الموظف"
    hdr[2].text = "عدد النشاطات"
    hdr[3].text = "القسم المنفّذ"
    hdr[4].text = "الأقسام المستهدفة"
    for c in hdr:
        for p in c.paragraphs:
            _set_rtl(p)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs:
                r.bold = True

    for i, row in enumerate(data, 1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = row["employee_name"]
        cells[2].text = str(row["total"])
        cells[3].text = row["employee_department_name"]
        cells[4].text = "، ".join(f"{t['name']} ({t['count']})" for t in row["targets"]) or "-"
        for c in cells:
            for p in c.paragraphs:
                _set_rtl(p)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    footer = doc.add_paragraph()
    _set_rtl(footer)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f = footer.add_run("GCAAN ACTIVT © 2026 — تصميم المهندس معاد كاظم")
    f.font.size = Pt(9)
    f.font.color.rgb = RGBColor(0x98, 0xA2, 0xB3)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"comprehensive-report-{datetime.now(timezone.utc).strftime('%Y%m%d')}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )

# ---------- Admin: System Info + Backup + Restore ----------
BACKUP_COLLECTIONS = ["users", "departments", "divisions", "activities"]

@api.get("/admin/system-info")
async def system_info(user=Depends(require_roles("general_manager"))):
    """Return storage stats (per collection) — GM only."""
    try:
        db_stats = await db.command("dbStats")
    except Exception:
        db_stats = {}
    coll_stats = []
    for name in BACKUP_COLLECTIONS:
        count = await db[name].count_documents({})
        size_bytes = 0
        storage_bytes = 0
        try:
            cs = await db.command("collStats", name)
            size_bytes = int(cs.get("size", 0))
            storage_bytes = int(cs.get("storageSize", 0))
        except Exception:
            pass
        coll_stats.append({
            "name": name,
            "count": count,
            "size_bytes": size_bytes,
            "storage_bytes": storage_bytes,
        })
    return {
        "database": os.environ["DB_NAME"],
        "total_data_bytes": int(db_stats.get("dataSize", 0)),
        "total_storage_bytes": int(db_stats.get("storageSize", 0)),
        "index_bytes": int(db_stats.get("indexSize", 0)),
        "objects": int(db_stats.get("objects", 0)),
        "collections": coll_stats,
        "generated_at": now_iso(),
    }


# ---------- Subscription ----------
class SubscriptionSet(BaseModel):
    start_date: Optional[str] = None  # YYYY-MM-DD
    end_date: Optional[str] = None    # YYYY-MM-DD
    duration_days: Optional[int] = None
    note: Optional[str] = None


def _compute_subscription_view(sub: dict) -> dict:
    """Compute derived subscription fields (days_remaining, expired, percent)."""
    start = sub.get("start_date")
    end = sub.get("end_date")
    today = datetime.now(timezone.utc).date()
    days_remaining = 0
    days_total = 0
    percent_used = 0
    expired = True
    if end:
        try:
            end_d = datetime.strptime(end, "%Y-%m-%d").date()
            delta = (end_d - today).days
            days_remaining = max(delta, 0)
            expired = delta < 0
        except Exception:
            pass
    if start and end:
        try:
            start_d = datetime.strptime(start, "%Y-%m-%d").date()
            end_d = datetime.strptime(end, "%Y-%m-%d").date()
            days_total = max((end_d - start_d).days, 0)
            elapsed = max((today - start_d).days, 0)
            if days_total > 0:
                percent_used = min(round((elapsed / days_total) * 100, 1), 100.0)
        except Exception:
            pass
    return {
        **sub,
        "days_remaining": days_remaining,
        "days_total": days_total,
        "percent_used": percent_used,
        "expired": expired,
    }


@api.get("/admin/subscription")
async def get_subscription(user=Depends(get_current_user)):
    """Any authenticated user can view the subscription status."""
    sub = await db.settings.find_one({"key": "subscription"}, {"_id": 0})
    if not sub:
        # default: 30 days from today, not yet configured
        today = datetime.now(timezone.utc).date()
        default = {
            "key": "subscription",
            "start_date": today.strftime("%Y-%m-%d"),
            "end_date": (today + timedelta(days=30)).strftime("%Y-%m-%d"),
            "note": "الاشتراك الافتراضي — يُرجى ضبط التاريخ الفعلي من قِبَل المدير العام",
            "configured": False,
            "updated_at": now_iso(),
        }
        return _compute_subscription_view(default)
    return _compute_subscription_view(sub)


@api.post("/admin/subscription")
async def set_subscription(
    payload: SubscriptionSet,
    user=Depends(require_roles("general_manager")),
):
    """Set/update the server subscription window — GM only."""
    today = datetime.now(timezone.utc).date()
    start_str = payload.start_date
    end_str = payload.end_date
    # Validate & normalize
    if not start_str:
        start_str = today.strftime("%Y-%m-%d")
    try:
        start_d = datetime.strptime(start_str, "%Y-%m-%d").date()
    except Exception:
        raise HTTPException(status_code=400, detail="تاريخ البداية غير صحيح (YYYY-MM-DD)")
    if end_str:
        try:
            end_d = datetime.strptime(end_str, "%Y-%m-%d").date()
        except Exception:
            raise HTTPException(status_code=400, detail="تاريخ الانتهاء غير صحيح (YYYY-MM-DD)")
    elif payload.duration_days is not None:
        if payload.duration_days < 1 or payload.duration_days > 3650:
            raise HTTPException(status_code=400, detail="مدة الاشتراك يجب أن تكون بين 1 و3650 يوماً")
        end_d = start_d + timedelta(days=payload.duration_days)
    else:
        raise HTTPException(status_code=400, detail="يجب تحديد تاريخ الانتهاء أو المدة")
    if end_d < start_d:
        raise HTTPException(status_code=400, detail="تاريخ الانتهاء يجب أن يكون بعد تاريخ البداية")

    doc = {
        "key": "subscription",
        "start_date": start_d.strftime("%Y-%m-%d"),
        "end_date": end_d.strftime("%Y-%m-%d"),
        "note": (payload.note or "").strip(),
        "configured": True,
        "updated_at": now_iso(),
        "updated_by": user["id"],
        "updated_by_name": user["full_name"],
    }
    await db.settings.update_one({"key": "subscription"}, {"$set": doc}, upsert=True)
    await log_audit(
        user,
        action="update",
        entity_type="subscription",
        details={"start_date": doc["start_date"], "end_date": doc["end_date"]},
    )
    saved = await db.settings.find_one({"key": "subscription"}, {"_id": 0})
    return _compute_subscription_view(saved)


@api.get("/admin/backup")
async def backup_download(user=Depends(require_roles("general_manager"))):
    """Download full JSON backup of all collections — GM only."""
    payload = {
        "version": 1,
        "generated_at": now_iso(),
        "database": os.environ["DB_NAME"],
        "collections": {},
    }
    for name in BACKUP_COLLECTIONS:
        docs = await db[name].find({}, {"_id": 0}).to_list(100000)
        payload["collections"][name] = docs
    buf = io.BytesIO(json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8"))
    buf.seek(0)
    filename = f"gcaan-backup-{datetime.now(timezone.utc).strftime('%Y%m%d-%H%M%S')}.json"
    await log_audit(user, action="backup", entity_type="system", details={"filename": filename})
    return StreamingResponse(
        buf,
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@api.post("/admin/restore")
async def restore_from_backup(
    payload: dict = Body(...),
    user=Depends(require_roles("general_manager")),
):
    """Restore full backup: replaces all collections. GM only."""
    collections = payload.get("collections")
    if not isinstance(collections, dict):
        raise HTTPException(status_code=400, detail="صيغة النسخة الاحتياطية غير صحيحة")
    # Basic validation: admin user must be present to keep GM login working
    users_backup = collections.get("users", [])
    if not any(u.get("role") == "general_manager" for u in users_backup):
        raise HTTPException(status_code=400, detail="النسخة الاحتياطية لا تحتوي على مدير عام. يُرفض للحفاظ على الأمان")
    restored: dict = {}
    for name in BACKUP_COLLECTIONS:
        docs = collections.get(name, []) or []
        await db[name].delete_many({})
        if docs:
            # deep copy to avoid mutating input
            clean_docs = [{k: v for k, v in d.items() if k != "_id"} for d in docs]
            await db[name].insert_many(clean_docs)
        restored[name] = len(docs)
    await log_audit(user, action="restore", entity_type="system", details={"restored": restored})
    return {"ok": True, "restored": restored, "at": now_iso()}


@api.get("/admin/audit-logs")
async def list_audit_logs(
    action: Optional[str] = Query(None),
    entity_type: Optional[str] = Query(None),
    user_id: Optional[str] = Query(None),
    date_from: Optional[str] = Query(None),
    date_to: Optional[str] = Query(None),
    limit: int = Query(200, ge=1, le=1000),
    user=Depends(require_roles("general_manager")),
):
    """List audit logs with filters — GM only."""
    q: dict = {}
    if action:
        q["action"] = action
    if entity_type:
        q["entity_type"] = entity_type
    if user_id:
        q["user_id"] = user_id
    if date_from or date_to:
        rng = {}
        if date_from:
            rng["$gte"] = date_from
        if date_to:
            rng["$lte"] = date_to
        q["at"] = rng
    docs = await db.audit_logs.find(q, {"_id": 0}).sort("at", -1).to_list(limit)
    total = await db.audit_logs.count_documents(q)
    return {"total": total, "count": len(docs), "logs": docs}


@api.get("/admin/audit-logs/summary")
async def audit_logs_summary(user=Depends(require_roles("general_manager"))):
    """Summary of audit logs — GM only."""
    total = await db.audit_logs.count_documents({})
    # Last 24h
    since_24h = (datetime.now(timezone.utc) - timedelta(hours=24)).isoformat()
    last_24h = await db.audit_logs.count_documents({"at": {"$gte": since_24h}})
    # Failed logins in last 24h
    failed_logins = await db.audit_logs.count_documents({
        "action": "login",
        "status": "failed",
        "at": {"$gte": since_24h},
    })
    # By action
    by_action_cur = db.audit_logs.aggregate([
        {"$group": {"_id": "$action", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}},
    ])
    by_action = [{"action": r["_id"], "count": r["count"]} async for r in by_action_cur]
    return {
        "total": total,
        "last_24h": last_24h,
        "failed_logins_24h": failed_logins,
        "by_action": by_action,
    }


@api.get("/admin/analysis-doc")
async def download_analysis_doc(user=Depends(require_roles("general_manager"))):
    """Download the comprehensive system analysis DOCX document."""
    doc_path = "/app/docs/GCAAN_ACTIVT_Analysis.docx"
    if not os.path.exists(doc_path):
        raise HTTPException(status_code=404, detail="ملف التحليل غير موجود")
    with open(doc_path, "rb") as f:
        data = f.read()
    buf = io.BytesIO(data)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="GCAAN_ACTIVT_Analysis.docx"'},
    )


# ---------- Mount ----------
app.include_router(api)
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)
